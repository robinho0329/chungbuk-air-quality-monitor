"""가설검정 리포트 자동 생성 (Markdown + Word).

측정소당 최소 표본(30건)이 모이면 산단 영향군 vs 베이스라인 t-test와
측정소 간 ANOVA를 오염물질별로 수행하고, 결과를 다음 두 곳에 남긴다:
  - reports/hypothesis/YYYY-MM-DD.md  (+ latest.md)
  - reports/hypothesis/YYYY-MM-DD.docx (Word)

표본 부족 시: 검정을 건너뛰고 진행 상황만 출력 (exit 0).
→ daily_dev_loop 워크플로우에 물려 매일 자동 실행되며, 조건 충족 시점부터 리포트 생성.

실행:
    uv run python scripts/generate_hypothesis_report.py
"""

from __future__ import annotations

import sys
from collections import Counter
from datetime import datetime, timedelta, timezone
from pathlib import Path

_PROJECT_ROOT = Path(__file__).resolve().parent.parent
if str(_PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(_PROJECT_ROOT))

import pandas as pd  # noqa: E402

from src.analysis.hypothesis_test import (  # noqa: E402
    MIN_SAMPLE_SIZE,
    AnovaResult,
    InsufficientSampleError,
    TTestResult,
    anova_across_stations,
    industrial_vs_baseline,
)
from src.analysis.usl_lsl import SPEC_LIMITS  # noqa: E402
from src.config import (  # noqa: E402
    BASELINE_GROUP,
    INDUSTRIAL_GROUP,
    STATION_GROUPS,
    TARGET_STATIONS,
)
from src.storage.database import query_all  # noqa: E402

KST = timezone(timedelta(hours=9))
REPORTS_DIR = _PROJECT_ROOT / "reports" / "hypothesis"
POLLUTANTS = list(SPEC_LIMITS.keys())  # pm10, pm25, o3, no2, so2, co
POLLUTANT_KR = {p: SPEC_LIMITS[p].description for p in POLLUTANTS}


def _load_df() -> pd.DataFrame:
    rows = query_all()
    if not rows:
        return pd.DataFrame()
    return pd.DataFrame.from_records(
        [
            {
                "station_name": r.station_name,
                "data_time": r.data_time,
                **{p: getattr(r, p) for p in POLLUTANTS},
            }
            for r in rows
        ]
    )


def _min_per_station(df: pd.DataFrame) -> int:
    if df.empty:
        return 0
    counts = Counter(df["station_name"])
    return min(counts.get(s, 0) for s in TARGET_STATIONS)


def run_tests(
    df: pd.DataFrame,
) -> tuple[dict[str, TTestResult], dict[str, AnovaResult]]:
    """오염물질별 t-test(산단 vs 베이스) + ANOVA(측정소 간) 수행."""
    ttests: dict[str, TTestResult] = {}
    anovas: dict[str, AnovaResult] = {}
    for p in POLLUTANTS:
        try:
            ttests[p] = industrial_vs_baseline(
                df, p, STATION_GROUPS, INDUSTRIAL_GROUP, BASELINE_GROUP
            )
        except (InsufficientSampleError, ValueError):
            pass
        try:
            anovas[p] = anova_across_stations(df, p)
        except (InsufficientSampleError, ValueError):
            pass
    return ttests, anovas


# ---------------------------------------------------------------------------
# Markdown 렌더
# ---------------------------------------------------------------------------
def render_markdown(
    today: datetime,
    df: pd.DataFrame,
    ttests: dict[str, TTestResult],
    anovas: dict[str, AnovaResult],
) -> str:
    L: list[str] = []
    L.append(f"# 🧪 가설검정 리포트 — {today.strftime('%Y-%m-%d')} (KST)")
    L.append("")
    L.append(
        f"> 자동 생성: {today.strftime('%Y-%m-%d %H:%M KST')} · "
        f"표본: 측정소당 최소 {_min_per_station(df)}건 / 총 {len(df)}건"
    )
    L.append("")
    L.append(
        "**검정 설계**: 산단 영향군(오창·복대·봉명·오송) vs 베이스라인(용암)을 "
        "Welch t-test로, 측정소 간 차이를 one-way ANOVA로 검정. "
        "p<0.05 유의, 효과크기(Cohen's d / η²) 병기. "
        "가설 정의는 `docs/ANALYSIS_HYPOTHESES.md` 참조."
    )
    L.append("")

    # H2/H3: t-test
    L.append("## H2·H3 — 산단 영향군 vs 베이스라인 (Welch t-test)")
    L.append("")
    L.append("| 오염물질 | 산단 평균 | 베이스 평균 | 차이 | p-value | 유의 | 효과크기(d) |")
    L.append("|---|---|---|---|---|---|---|")
    for p in POLLUTANTS:
        r = ttests.get(p)
        if r is None:
            L.append(f"| {POLLUTANT_KR[p]} | — | — | — | — | 표본부족 | — |")
            continue
        sig = "✅ 유의" if r.significant else "✖ 비유의"
        L.append(
            f"| {POLLUTANT_KR[p]} | {r.mean_a:.3f} | {r.mean_b:.3f} | "
            f"{r.diff:+.3f} | {r.p_value:.4f} | {sig} | "
            f"{r.cohens_d:+.2f} ({r.effect_label()}) |"
        )
    L.append("")
    for p in POLLUTANTS:
        r = ttests.get(p)
        if r is not None:
            L.append(f"- **{POLLUTANT_KR[p]}**: {r.interpret()}")
    L.append("")

    # ANOVA
    L.append("## 측정소 간 차이 (one-way ANOVA)")
    L.append("")
    L.append("| 오염물질 | F | p-value | 유의 | η² | 최고 측정소 |")
    L.append("|---|---|---|---|---|---|")
    for p in POLLUTANTS:
        r = anovas.get(p)
        if r is None:
            L.append(f"| {POLLUTANT_KR[p]} | — | — | 표본부족 | — | — |")
            continue
        sig = "✅ 유의" if r.significant else "✖ 비유의"
        hi = r.labels[r.group_means.index(max(r.group_means))]
        L.append(
            f"| {POLLUTANT_KR[p]} | {r.f_stat:.2f} | {r.p_value:.4f} | "
            f"{sig} | {r.eta_squared:.3f} | {hi} |"
        )
    L.append("")

    # 종합 해석
    L.append("## 종합 해석 (DMAIC Analyze)")
    L.append("")
    pm25 = ttests.get("pm25")
    if pm25 is not None:
        if pm25.significant and pm25.diff > 0:
            L.append(
                "- **H3 기각 / H2 지지**: 산단 영향군의 PM2.5가 베이스라인보다 "
                f"통계적으로 유의하게 높음(p={pm25.p_value:.4f}). "
                "산단 인접 지역의 초미세먼지 부하가 거주지보다 크다는 가설을 지지."
            )
        elif not pm25.significant:
            L.append(
                "- **H3 지지 가능**: 산단 영향군과 베이스라인의 PM2.5 차이가 "
                "통계적으로 유의하지 않음 → 산단 자체보다 광역 기상·교통이 "
                "지배적일 가능성. (저배출 업종 가설과 부합)"
            )
    L.append(
        "- ⚠️ 본 검정은 누적 데이터 기준 단면 분석이다. 평일/주말 분리와 "
        "다일 누적, 자기상관(시계열) 보정은 후속 과제(WEATHER_API·다변량)."
    )
    L.append("")
    L.append("---")
    L.append("🤖 *데이터가 누적될수록 매일 자동 재검정됩니다.*")
    return "\n".join(L)


# ---------------------------------------------------------------------------
# Word(docx) 렌더
# ---------------------------------------------------------------------------
def render_docx(
    today: datetime,
    df: pd.DataFrame,
    ttests: dict[str, TTestResult],
    anovas: dict[str, AnovaResult],
    out_path: Path,
) -> None:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading("충북권 산업단지 대기질 — 가설검정 리포트", level=0)
    p = doc.add_paragraph()
    p.add_run(
        f"생성일: {today.strftime('%Y-%m-%d %H:%M KST')}  |  "
        f"표본: 측정소당 최소 {_min_per_station(df)}건 / 총 {len(df)}건"
    ).italic = True

    doc.add_heading("1. 검정 설계", level=1)
    doc.add_paragraph(
        "산단 영향군(오창·복대·봉명·오송) vs 베이스라인(용암)을 Welch t-test로, "
        "측정소 간 차이를 one-way ANOVA로 검정한다. 유의수준 0.05, "
        "효과크기(Cohen's d / η²)를 병기하여 통계적·실질적 유의성을 함께 평가한다."
    )

    # t-test 표
    doc.add_heading("2. 산단 영향군 vs 베이스라인 (Welch t-test)", level=1)
    headers = ["오염물질", "산단 평균", "베이스 평균", "차이", "p-value", "유의", "효과크기 d"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for pol in POLLUTANTS:
        r = ttests.get(pol)
        cells = table.add_row().cells
        if r is None:
            cells[0].text = POLLUTANT_KR[pol]
            for j in range(1, len(headers)):
                cells[j].text = "표본부족" if j == 5 else "—"
            continue
        vals = [
            POLLUTANT_KR[pol],
            f"{r.mean_a:.3f}",
            f"{r.mean_b:.3f}",
            f"{r.diff:+.3f}",
            f"{r.p_value:.4f}",
            "유의" if r.significant else "비유의",
            f"{r.cohens_d:+.2f} ({r.effect_label()})",
        ]
        for j, v in enumerate(vals):
            cells[j].text = v

    # ANOVA 표
    doc.add_heading("3. 측정소 간 차이 (one-way ANOVA)", level=1)
    headers2 = ["오염물질", "F", "p-value", "유의", "η²", "최고 측정소"]
    table2 = doc.add_table(rows=1, cols=len(headers2))
    table2.style = "Light Grid Accent 1"
    for i, h in enumerate(headers2):
        table2.rows[0].cells[i].text = h
    for pol in POLLUTANTS:
        r = anovas.get(pol)
        cells = table2.add_row().cells
        if r is None:
            cells[0].text = POLLUTANT_KR[pol]
            for j in range(1, len(headers2)):
                cells[j].text = "표본부족" if j == 3 else "—"
            continue
        hi = r.labels[r.group_means.index(max(r.group_means))]
        vals = [
            POLLUTANT_KR[pol],
            f"{r.f_stat:.2f}",
            f"{r.p_value:.4f}",
            "유의" if r.significant else "비유의",
            f"{r.eta_squared:.3f}",
            hi,
        ]
        for j, v in enumerate(vals):
            cells[j].text = v

    # 해석
    doc.add_heading("4. 종합 해석 (DMAIC Analyze)", level=1)
    pm25 = ttests.get("pm25")
    if pm25 is not None:
        doc.add_paragraph(pm25.interpret(), style="List Bullet")
    doc.add_paragraph(
        "본 검정은 누적 데이터 단면 분석이며, 평일/주말 분리·다일 누적·"
        "시계열 자기상관 보정은 후속 과제다.",
        style="List Bullet",
    )

    note = doc.add_paragraph()
    run = note.add_run("자동 생성 — 데이터 누적 시 매일 재검정됨.")
    run.font.size = Pt(8)
    run.italic = True

    doc.save(str(out_path))


# ---------------------------------------------------------------------------
# 메인
# ---------------------------------------------------------------------------
def main() -> int:
    today = datetime.now(tz=KST)
    df = _load_df()
    min_per = _min_per_station(df)
    print(f"=== 가설검정 리포트 @ {today.strftime('%Y-%m-%d %H:%M KST')} ===")
    print(f"총 {len(df)}건 / 측정소당 최소 {min_per}건 (요건 {MIN_SAMPLE_SIZE})")

    if min_per < MIN_SAMPLE_SIZE:
        remaining = MIN_SAMPLE_SIZE - min_per
        print(
            f"⏳ 표본 부족 — 측정소당 {remaining}건 더 필요. "
            "리포트 생성을 건너뜁니다 (데이터 누적 후 자동 실행)."
        )
        return 0

    ttests, anovas = run_tests(df)
    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    stamp = today.strftime("%Y-%m-%d")

    md = render_markdown(today, df, ttests, anovas)
    (REPORTS_DIR / f"{stamp}.md").write_text(md, encoding="utf-8")
    (REPORTS_DIR / "latest.md").write_text(md, encoding="utf-8")
    render_docx(today, df, ttests, anovas, REPORTS_DIR / f"{stamp}.docx")

    print(f"✅ 리포트 생성: reports/hypothesis/{stamp}.md / .docx")
    print(f"   t-test {len(ttests)}건, ANOVA {len(anovas)}건 수행")
    return 0


if __name__ == "__main__":
    sys.exit(main())
